from __future__ import annotations

import json
import sys
import zipfile
from datetime import date
from pathlib import Path

import joblib
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[1]
BACKEND_DIR = ROOT_DIR / "backend"
REPORTS_DIR = ROOT_DIR / "reports"
FIGURES_DIR = REPORTS_DIR / "figures"
OUTPUT_PATH = REPORTS_DIR / "Fish_Kiln_Beginner_Results_Explanation.docx"

sys.path.insert(0, str(BACKEND_DIR))
sys.path.insert(0, str(ROOT_DIR / "scripts"))

from app import build_prediction_features, drying_status  # noqa: E402
from generate_results_report import (  # noqa: E402
    TEST_BATCH,
    clean_feature_name,
    evaluate_model,
    load_training_frame,
    plot_decision_matrix,
    run_app_checks,
)
from train_model import (  # noqa: E402
    DATA_PATH,
    FEATURE_COLUMNS,
    FISH_MOISTURE_CONTENT,
    METADATA_PATH,
    MODEL_PATH,
    TARGET_MOISTURE_CONTENT,
)


ACCENT = RGBColor(0, 127, 115)
INK = RGBColor(23, 32, 28)
MUTED = RGBColor(101, 113, 107)


def fix_docx_zoom_setting(path: Path) -> None:
    tmp_path = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/settings.xml":
                text = data.decode("utf-8").replace(
                    '<w:zoom w:val="bestFit"/>',
                    '<w:zoom w:val="bestFit" w:percent="100"/>',
                )
                data = text.encode("utf-8")
            target.writestr(item, data)
    tmp_path.replace(path)


def add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = INK

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED


def add_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Inches(0.2)
    paragraph_format.right_indent = Inches(0.2)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = ACCENT


def add_table(document: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_picture_if_exists(document: Document, title: str, path: Path, explanation: str) -> None:
    if not path.exists():
        return
    document.add_heading(title, level=2)
    document.add_picture(str(path), width=Inches(6.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(explanation)


def main() -> None:
    REPORTS_DIR.mkdir(exist_ok=True)
    data = load_training_frame()
    evaluation = evaluate_model(data)
    metrics = evaluation["metrics"]
    _, decision_matrix = plot_decision_matrix(evaluation["y_test"], evaluation["predictions"])
    app_checks, demo_prediction = run_app_checks()
    metadata = json.loads(METADATA_PATH.read_text(encoding="utf-8"))

    model = joblib.load(MODEL_PATH)
    feature_names = [clean_feature_name(name) for name in model.named_steps["preprocessor"].get_feature_names_out()]
    importances = (
        pd.Series(model.named_steps["regressor"].feature_importances_, index=feature_names)
        .sort_values(ascending=False)
        .round(4)
    )

    species_counts = data["fish_type"].value_counts().sort_index()
    x_train_rows = int(len(data) * 0.8)
    x_test_rows = len(data) - x_train_rows
    app_features = build_prediction_features(
        fish_type=TEST_BATCH["fishType"],
        initial_weight=TEST_BATCH["initialWeight"],
        current_weight=TEST_BATCH["currentWeight"],
        temperature=TEST_BATCH["temperature"],
        humidity=TEST_BATCH["humidity"],
        elapsed_time=TEST_BATCH["elapsedTime"],
    )

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = INK

    add_title(
        document,
        "Beginner Explanation of the Fish Kiln AI Results",
        "A viva/supervisor guide for explaining where each result came from",
    )
    document.add_paragraph(f"Generated: {date.today().isoformat()}")

    document.add_heading("1. The Short Answer You Can Say First", level=1)
    document.add_paragraph(
        "The application predicts the remaining drying time of fish in a kiln. It takes fish type, weight readings, temperature, humidity, and elapsed drying time, then calculates moisture-related values. A Random Forest Regression model uses those values to estimate how many minutes are left before the fish reaches the target moisture content."
    )
    add_note(
        document,
        "Be honest: the current model is trained on synthetic demonstration data generated by train_model.py. The next academic improvement is to collect real kiln experiment data and retrain the model.",
    )

    document.add_heading("2. Where Did the Dataset Come From?", level=1)
    document.add_paragraph(
        "The dataset used in the current project is not downloaded from the internet and it is not yet real laboratory/field data. It is generated locally by the function generate_synthetic_data() in backend/train_model.py. The file produced is backend/synthetic_drying_data.csv."
    )
    add_table(
        document,
        ["Question", "Answer you should give"],
        [
            ["Dataset source", "Generated locally from backend/train_model.py"],
            ["Dataset file", str(DATA_PATH.relative_to(ROOT_DIR))],
            ["Dataset type", metadata.get("trainingDataType", "synthetic")],
            ["Rows", len(data)],
            ["Rows per species", "; ".join(f"{species}: {count}" for species, count in species_counts.items())],
            ["Target variable", metadata.get("target", "remaining_drying_time")],
            ["Model type", metadata.get("modelType", "Random Forest Regression")],
        ],
    )

    document.add_paragraph(
        "Why synthetic data was used: the app needed a working model before real ESP32/kiln drying trials were available. The synthetic data allows the software pipeline to be built, tested, deployed, and demonstrated. But it should not be claimed as real experimental evidence."
    )

    document.add_heading("3. What Each Dataset Column Means", level=1)
    add_table(
        document,
        ["Column", "Meaning"],
        [
            ["fish_type", "Species category: catfish, tilapia, or mackerel."],
            ["initial_weight", "Starting fish weight before drying, in kg."],
            ["current_weight", "Fish weight at the current drying time, in kg."],
            ["temperature", "Kiln temperature assumption, in C."],
            ["humidity", "Humidity assumption, in percent."],
            ["elapsed_time", "How long drying has already been running, in minutes."],
            ["initial_moisture_content", "Starting moisture percentage assumed for that fish species."],
            ["estimated_moisture_content", "Calculated current moisture percentage."],
            ["weight_loss", "Initial weight minus current weight."],
            ["percentage_weight_loss", "Weight loss as a percentage of the initial weight."],
            ["drying_progress", "How far the fish has moved from initial moisture toward the 15% target."],
            ["remaining_drying_time", "The value the model learns to predict, in minutes."],
        ],
    )

    document.add_heading("4. Where the Main Assumptions Came From", level=1)
    add_table(
        document,
        ["Item", "Value in the code", "Purpose"],
        [
            ["Target moisture content", f"{TARGET_MOISTURE_CONTENT}%", "The drying goal used by both backend/app.py and train_model.py."],
            ["Catfish initial moisture", f"{FISH_MOISTURE_CONTENT['catfish']}%", "Starting moisture assumption for catfish."],
            ["Tilapia initial moisture", f"{FISH_MOISTURE_CONTENT['tilapia']}%", "Starting moisture assumption for tilapia."],
            ["Mackerel initial moisture", f"{FISH_MOISTURE_CONTENT['mackerel']}%", "Starting moisture assumption for mackerel."],
            ["Initial weight range", "0.5 kg to 5.0 kg", "Randomly sampled in train_model.py."],
            ["Temperature range", "45 C to 70 C", "Randomly sampled in train_model.py."],
            ["Humidity range", "25% to 85%", "Randomly sampled in train_model.py."],
            ["Elapsed time range", "0 to 360 minutes", "Randomly sampled in train_model.py."],
        ],
    )
    document.add_paragraph(
        "These assumptions are encoded in the Python script, so if you are asked 'where did you get it?', the accurate answer is: they are project assumptions used to simulate drying behavior until real data is collected. They should be replaced or calibrated with real experimental values later."
    )

    document.add_heading("5. How the Synthetic Rows Were Created", level=1)
    document.add_paragraph(
        "For each fish species, the script creates 700 rows. Since there are three species, the total becomes 700 x 3 = 2100 rows. A fixed random_state of 42 is used so the same data can be regenerated again."
    )
    add_numbered(
        document,
        [
            "Pick a fish species: catfish, tilapia, or mackerel.",
            "Randomly choose an initial weight, temperature, humidity, and elapsed time.",
            "Use species drying assumptions to estimate a total drying time.",
            "Estimate how far drying has progressed at the chosen elapsed time.",
            "Estimate the current moisture content.",
            "Convert moisture content into current weight.",
            "Calculate weight loss, percentage weight loss, and drying progress.",
            "Generate remaining_drying_time, which becomes the target value for the ML model.",
        ],
    )

    document.add_heading("6. The Key Calculations in Simple English", level=1)
    document.add_paragraph("The backend uses drying logic before the ML model is called. These are the formulas you should understand:")
    add_table(
        document,
        ["Calculation", "Formula", "Beginner meaning"],
        [
            ["Dry matter", "initial_weight x (1 - initial_moisture/100)", "The part of the fish that is not water."],
            ["Estimated moisture", "((current_weight - dry_matter) / current_weight) x 100", "How much water is still inside the fish by percentage."],
            ["Weight loss", "initial_weight - current_weight", "How much mass has been lost during drying."],
            ["Percentage weight loss", "(weight_loss / initial_weight) x 100", "Weight loss expressed as a percentage."],
            ["Drying progress", "((initial_moisture - estimated_moisture) / (initial_moisture - target_moisture)) x 100", "How close the fish is to the 15% target moisture."],
        ],
    )

    document.add_heading("7. How the Model Was Trained", level=1)
    document.add_paragraph(
        "The model is trained in backend/train_model.py. The input features are stored in FEATURE_COLUMNS. The output/target is remaining_drying_time."
    )
    add_table(
        document,
        ["Training step", "What happened"],
        [
            ["Input X", ", ".join(FEATURE_COLUMNS)],
            ["Target y", "remaining_drying_time"],
            ["Train/test split", f"80% train ({x_train_rows} rows), 20% test ({x_test_rows} rows)"],
            ["Categorical preprocessing", "fish_type is converted with OneHotEncoder."],
            ["Numeric preprocessing", "Numeric columns are scaled with StandardScaler."],
            ["Model", "RandomForestRegressor with 240 trees and min_samples_leaf=2."],
            ["Saved model", str(MODEL_PATH.relative_to(ROOT_DIR))],
            ["Saved metrics", str(METADATA_PATH.relative_to(ROOT_DIR))],
        ],
    )
    document.add_paragraph(
        "A Random Forest is an ensemble model. That means it builds many decision trees and averages their outputs. It is useful here because the relationship between moisture, temperature, humidity, weight loss, and remaining time is not perfectly linear."
    )

    document.add_heading("8. How You Arrived at MAE, RMSE, and R2", level=1)
    document.add_paragraph(
        "After training, the model was tested on the 20% test set. These rows were held back from training. The model predicted remaining drying time for those rows, then the predictions were compared with the known target values."
    )
    add_table(
        document,
        ["Metric", "Your result", "How to explain it"],
        [
            ["MAE", metrics["mae"], "Mean Absolute Error. On average, predictions are about 14.222 minutes away from the synthetic target."],
            ["RMSE", metrics["rmse"], "Root Mean Squared Error. It also measures error in minutes but punishes large mistakes more strongly."],
            ["R2", metrics["r2"], "The model explains about 97.9% of the variation in the synthetic test data."],
        ],
    )
    add_note(
        document,
        "Do not oversell the R2. It is high mainly because the data was generated by a controlled synthetic formula. Real kiln data will probably be noisier, and the R2 may reduce after real-world retraining.",
    )

    document.add_heading("9. Why There Is No Normal Confusion Matrix", level=1)
    document.add_paragraph(
        "A confusion matrix is normally used for classification problems, for example predicting 'dry' or 'not dry'. Your model is not doing classification. It predicts a continuous number: remaining drying time in minutes. That makes it a regression problem."
    )
    document.add_paragraph(
        "To still provide something similar for presentation, I converted the continuous predicted time into operational categories:"
    )
    add_table(
        document,
        ["Category", "Rule"],
        [
            ["Ready/complete", "0 to 10 minutes remaining"],
            ["Short wait", "More than 10 and up to 90 minutes remaining"],
            ["Needs drying", "More than 90 minutes remaining"],
        ],
    )
    document.add_paragraph("The resulting operational decision matrix on the 420 test rows was:")
    add_table(
        document,
        ["Actual / Predicted", *[str(column) for column in decision_matrix.columns]],
        [[str(index), *row.tolist()] for index, row in decision_matrix.iterrows()],
    )

    document.add_heading("10. What the Graphs Mean", level=1)
    add_picture_if_exists(
        document,
        "Feature Importance",
        FIGURES_DIR / "feature_importance.png",
        "The most important feature is estimated moisture content. This makes sense because remaining drying time is strongly controlled by how much moisture is still left in the fish.",
    )
    add_table(
        document,
        ["Feature", "Importance"],
        [[name, value] for name, value in importances.head(8).items()],
    )
    add_picture_if_exists(
        document,
        "Actual vs Predicted",
        FIGURES_DIR / "actual_vs_predicted.png",
        "If the model is accurate, the points stay close to the diagonal line. Points far from the line are larger errors.",
    )
    add_picture_if_exists(
        document,
        "Residual Distribution",
        FIGURES_DIR / "residual_distribution.png",
        "Residual means actual value minus predicted value. A good model should have most residuals near zero.",
    )
    add_picture_if_exists(
        document,
        "Correlation Heatmap",
        FIGURES_DIR / "correlation_heatmap.png",
        "The heatmap shows how numeric columns move together. Strong positive or negative correlations help explain relationships in the drying data.",
    )

    document.add_heading("11. Explain the Demo Prediction Step by Step", level=1)
    document.add_paragraph("The demo/test batch used by the app is:")
    add_table(
        document,
        ["Input", "Value"],
        [[key, value] for key, value in TEST_BATCH.items()],
    )
    document.add_paragraph("For catfish, the assumed initial moisture is 65%. The target moisture is 15%.")
    add_table(
        document,
        ["Step", "Calculation", "Result"],
        [
            ["Dry matter", "2.0 x (1 - 65/100)", f"{app_features['dry_matter']:.3f} kg"],
            ["Weight loss", "2.0 - 1.4", f"{app_features['weight_loss']:.3f} kg"],
            ["Percentage weight loss", "(0.6 / 2.0) x 100", f"{app_features['percentage_weight_loss']:.2f}%"],
            ["Estimated moisture", "((1.4 - 0.7) / 1.4) x 100", f"{app_features['estimated_moisture_content']:.2f}%"],
            ["Drying progress", "((65 - 50) / (65 - 15)) x 100", f"{app_features['drying_progress']:.2f}%"],
            ["Drying status rule", "50% is above 25%", drying_status(app_features["estimated_moisture_content"], False)],
            ["Model output", "Random Forest prediction", f"{demo_prediction.get('estimatedRemainingTime')} minutes remaining"],
        ],
    )
    document.add_paragraph(
        "So, for that demo case, the fish is still drying. It has lost 0.6 kg, estimated moisture is 50%, drying progress is 30%, and the model estimates about 218.67 minutes remaining."
    )

    document.add_heading("12. What the App Validation Matrix Means", level=1)
    document.add_paragraph(
        "The app validation matrix is not an AI metric. It is a software testing matrix. It shows that the frontend builds correctly, the backend endpoints respond, prediction works, and invalid inputs are rejected."
    )
    add_table(
        document,
        ["Area", "Check", "Status", "Evidence"],
        [[row["area"], row["check"], row["status"], row["evidence"]] for row in app_checks],
    )

    document.add_heading("13. Questions Your Supervisor May Ask", level=1)
    add_table(
        document,
        ["Likely question", "Answer"],
        [
            [
                "Where did you get the dataset?",
                "For the current version, I generated a synthetic dataset using train_model.py. It has 2100 rows and simulates fish drying behavior using moisture, weight, temperature, humidity, and elapsed time assumptions.",
            ],
            [
                "Is it real experimental data?",
                "No. It is demonstration data. The next step is to collect real kiln/ESP32 experimental data and retrain the model.",
            ],
            [
                "What exactly is the AI predicting?",
                "It predicts remaining_drying_time, which is the estimated number of minutes left before the fish reaches the target moisture level.",
            ],
            [
                "Why Random Forest?",
                "Because drying is affected by several non-linear factors. Random Forest handles non-linear relationships and mixed features well for this version of the system.",
            ],
            [
                "Why is R2 so high?",
                "Because the data is synthetic and generated from a controlled formula. It proves the workflow, but real data will be needed for a stronger scientific claim.",
            ],
            [
                "Why no normal confusion matrix?",
                "The model is regression, not classification. I used MAE, RMSE, R2, residual plots, and a binned operational decision matrix instead.",
            ],
            [
                "What does the app prove?",
                "It proves the workflow: collect input, calculate drying features, call the trained model, return prediction, show status, and validate bad inputs.",
            ],
        ],
    )

    document.add_heading("14. The Clean Way to Present the Limitation", level=1)
    document.add_paragraph(
        "A good way to say it is: 'At this stage, the model is trained on synthetic drying data generated from domain assumptions. The purpose is to demonstrate the intelligent monitoring workflow. For final validation, I will need to collect real drying readings from the kiln, retrain the model, and compare the model performance on real experimental test data.'"
    )

    document.add_heading("15. What to Improve Next", level=1)
    add_bullets(
        document,
        [
            "Collect real sensor readings: fish type, initial weight, current weight, temperature, humidity, elapsed time, and final moisture measurement.",
            "Start with one species, such as catfish, if collecting all species is too difficult.",
            "Retrain the model on real experimental data.",
            "Compare synthetic-trained results with real-data-trained results.",
            "Add a true classification model later if the supervisor specifically wants classes like Dry, Almost dry, and Not dry.",
            "Document the experimental procedure so the dataset has academic credibility.",
        ],
    )

    document.add_heading("16. One-Minute Presentation Script", level=1)
    document.add_paragraph(
        "This system estimates the remaining drying time for fish in a kiln. The current version uses a synthetic dataset generated in Python because real experiment data has not yet been collected. The dataset contains 2100 rows across catfish, tilapia, and mackerel. From weight, moisture, temperature, humidity, and elapsed time, the system calculates drying features such as weight loss and drying progress. A Random Forest Regression model was trained to predict remaining drying time. On the held-out synthetic test set, the model achieved MAE of 14.222 minutes, RMSE of 19.622 minutes, and R2 of 0.979. Since this is regression, I used regression metrics and diagnostic plots instead of a normal confusion matrix. The deployed app validates the workflow by accepting batch readings, calculating moisture state, calling the model, and returning drying status and remaining time. The major next step is to collect real kiln drying data and retrain the model for final validation."
    )

    document.save(OUTPUT_PATH)
    fix_docx_zoom_setting(OUTPUT_PATH)
    print(f"Explanation document written to: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
