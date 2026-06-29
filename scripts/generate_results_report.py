from __future__ import annotations

import json
import subprocess
import sys
import zipfile
from datetime import date
from pathlib import Path

import joblib
import matplotlib
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.model_selection import train_test_split


matplotlib.use("Agg")
from matplotlib import pyplot as plt  # noqa: E402

ROOT_DIR = Path(__file__).resolve().parents[1]
BACKEND_DIR = ROOT_DIR / "backend"
FRONTEND_DIR = ROOT_DIR / "frontend"
REPORTS_DIR = ROOT_DIR / "reports"
FIGURES_DIR = REPORTS_DIR / "figures"
REPORT_PATH = REPORTS_DIR / "Fish_Kiln_AI_Model_and_App_Results.docx"

sys.path.insert(0, str(BACKEND_DIR))

from app import app as flask_app  # noqa: E402
from train_model import (  # noqa: E402
    DATA_PATH,
    FEATURE_COLUMNS,
    METADATA_PATH,
    MODEL_PATH,
    generate_synthetic_data,
    train_model,
)


TEST_BATCH = {
    "fishType": "catfish",
    "initialWeight": 2.0,
    "currentWeight": 1.4,
    "temperature": 60,
    "humidity": 45,
    "elapsedTime": 120,
}

ACCENT = "#007f73"
BLUE = "#2563eb"
AMBER = "#b55a00"
INK = "#17201c"
MUTED = "#65716b"
LINE = "#d9e0da"


def ensure_artifacts() -> None:
    REPORTS_DIR.mkdir(exist_ok=True)
    FIGURES_DIR.mkdir(exist_ok=True)

    if not MODEL_PATH.exists() or not METADATA_PATH.exists():
        train_model()


def load_training_frame() -> pd.DataFrame:
    if DATA_PATH.exists():
        return pd.read_csv(DATA_PATH)
    return generate_synthetic_data()


def evaluate_model(data: pd.DataFrame) -> dict[str, object]:
    model = joblib.load(MODEL_PATH)
    x = data[FEATURE_COLUMNS]
    y = data["remaining_drying_time"]

    _, x_test, _, y_test = train_test_split(
        x,
        y,
        test_size=0.2,
        random_state=42,
    )
    predictions = model.predict(x_test)
    residuals = y_test.to_numpy() - predictions

    metrics = {
        "mae": round(float(mean_absolute_error(y_test, predictions)), 3),
        "rmse": round(float(np.sqrt(mean_squared_error(y_test, predictions))), 3),
        "r2": round(float(r2_score(y_test, predictions)), 3),
    }

    return {
        "model": model,
        "x_test": x_test,
        "y_test": y_test,
        "predictions": predictions,
        "residuals": residuals,
        "metrics": metrics,
    }


def clean_feature_name(name: str) -> str:
    return (
        name.replace("fish_type__fish_type_", "fish: ")
        .replace("numeric__", "")
        .replace("_", " ")
    )


def plot_correlation_heatmap(data: pd.DataFrame) -> Path:
    numeric_data = data.select_dtypes(include=["number"])
    corr = numeric_data.corr()

    fig, ax = plt.subplots(figsize=(10.5, 8.0))
    im = ax.imshow(corr, cmap="BrBG", vmin=-1, vmax=1)
    ax.set_title("Feature Correlation Heatmap", fontsize=15, fontweight="bold", color=INK)
    ax.set_xticks(range(len(corr.columns)))
    ax.set_yticks(range(len(corr.columns)))
    labels = [column.replace("_", " ") for column in corr.columns]
    ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=8)
    ax.set_yticklabels(labels, fontsize=8)
    for i in range(len(corr.columns)):
        for j in range(len(corr.columns)):
            ax.text(j, i, f"{corr.iloc[i, j]:.2f}", ha="center", va="center", fontsize=6)
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    path = FIGURES_DIR / "correlation_heatmap.png"
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_actual_vs_predicted(y_test: pd.Series, predictions: np.ndarray) -> Path:
    fig, ax = plt.subplots(figsize=(8.4, 5.8))
    ax.scatter(y_test, predictions, s=24, color=BLUE, alpha=0.64, edgecolors="none")
    low = min(float(y_test.min()), float(predictions.min()))
    high = max(float(y_test.max()), float(predictions.max()))
    ax.plot([low, high], [low, high], color=AMBER, linewidth=2.2, label="Perfect prediction")
    ax.set_title("Actual vs Predicted Remaining Drying Time", fontsize=15, fontweight="bold", color=INK)
    ax.set_xlabel("Actual remaining time (minutes)")
    ax.set_ylabel("Predicted remaining time (minutes)")
    ax.grid(True, color=LINE, linewidth=0.8)
    ax.legend()
    fig.tight_layout()
    path = FIGURES_DIR / "actual_vs_predicted.png"
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_residuals(residuals: np.ndarray) -> Path:
    fig, ax = plt.subplots(figsize=(8.4, 5.2))
    ax.hist(residuals, bins=28, color=ACCENT, alpha=0.86, edgecolor="white")
    ax.axvline(0, color=AMBER, linewidth=2.2)
    ax.set_title("Residual Error Distribution", fontsize=15, fontweight="bold", color=INK)
    ax.set_xlabel("Actual minus predicted remaining time (minutes)")
    ax.set_ylabel("Frequency")
    ax.grid(True, axis="y", color=LINE, linewidth=0.8)
    fig.tight_layout()
    path = FIGURES_DIR / "residual_distribution.png"
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def plot_feature_importance(model) -> Path:
    preprocessor = model.named_steps["preprocessor"]
    regressor = model.named_steps["regressor"]
    feature_names = [clean_feature_name(name) for name in preprocessor.get_feature_names_out()]
    importances = pd.Series(regressor.feature_importances_, index=feature_names).sort_values(ascending=True)
    top_features = importances.tail(10)

    fig, ax = plt.subplots(figsize=(8.6, 5.8))
    ax.barh(top_features.index, top_features.values, color=ACCENT)
    ax.set_title("Top Feature Importance Scores", fontsize=15, fontweight="bold", color=INK)
    ax.set_xlabel("Importance")
    ax.grid(True, axis="x", color=LINE, linewidth=0.8)
    fig.tight_layout()
    path = FIGURES_DIR / "feature_importance.png"
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def remaining_time_bucket(values: pd.Series | np.ndarray) -> pd.Categorical:
    labels = ["Ready/complete", "Short wait", "Needs drying"]
    return pd.cut(values, bins=[-0.01, 10, 90, np.inf], labels=labels)


def plot_decision_matrix(y_test: pd.Series, predictions: np.ndarray) -> tuple[Path, pd.DataFrame]:
    actual_bucket = remaining_time_bucket(y_test.to_numpy())
    predicted_bucket = remaining_time_bucket(predictions)
    matrix = pd.crosstab(
        pd.Series(actual_bucket, name="Actual"),
        pd.Series(predicted_bucket, name="Predicted"),
        dropna=False,
    )

    fig, ax = plt.subplots(figsize=(7.8, 5.6))
    im = ax.imshow(matrix.to_numpy(), cmap="YlGnBu")
    ax.set_title("Operational Decision Matrix", fontsize=15, fontweight="bold", color=INK)
    ax.set_xticks(range(len(matrix.columns)))
    ax.set_yticks(range(len(matrix.index)))
    ax.set_xticklabels(matrix.columns, rotation=20, ha="right")
    ax.set_yticklabels(matrix.index)
    ax.set_xlabel("Predicted class from model output")
    ax.set_ylabel("Actual class from test target")
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            ax.text(j, i, str(matrix.iloc[i, j]), ha="center", va="center", fontsize=12, fontweight="bold")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    path = FIGURES_DIR / "operational_decision_matrix.png"
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path, matrix


def run_app_checks() -> tuple[list[dict[str, str]], dict[str, object]]:
    rows: list[dict[str, str]] = []
    prediction_payload: dict[str, object] = {}

    build = subprocess.run(
        ["npm", "run", "build"],
        cwd=FRONTEND_DIR,
        text=True,
        capture_output=True,
        check=False,
    )
    rows.append(
        {
            "area": "Frontend",
            "check": "Production build",
            "status": "Pass" if build.returncode == 0 else "Fail",
            "evidence": "Vite build completed" if build.returncode == 0 else build.stderr[-160:],
        }
    )

    assets_dir = FRONTEND_DIR / "dist" / "assets"
    if assets_dir.exists():
        total_kb = sum(path.stat().st_size for path in assets_dir.iterdir() if path.is_file()) / 1024
        rows.append(
            {
                "area": "Frontend",
                "check": "Compiled asset size",
                "status": "Recorded",
                "evidence": f"{total_kb:.1f} KB in dist/assets",
            }
        )

    with flask_app.test_client() as client:
        health = client.get("/api/health")
        rows.append(
            {
                "area": "Backend API",
                "check": "Health endpoint",
                "status": "Pass" if health.status_code == 200 else "Fail",
                "evidence": f"HTTP {health.status_code}",
            }
        )

        model_info = client.get("/api/model-info")
        rows.append(
            {
                "area": "Backend API",
                "check": "Model metadata endpoint",
                "status": "Pass" if model_info.status_code == 200 else "Fail",
                "evidence": f"HTTP {model_info.status_code}",
            }
        )

        prediction = client.post("/api/predict", json=TEST_BATCH)
        prediction_payload = prediction.get_json() or {}
        rows.append(
            {
                "area": "Backend API",
                "check": "Prediction endpoint",
                "status": "Pass" if prediction.status_code == 200 else "Fail",
                "evidence": (
                    f"HTTP {prediction.status_code}; "
                    f"status={prediction_payload.get('dryingStatus', 'n/a')}; "
                    f"remaining={prediction_payload.get('estimatedRemainingTime', 'n/a')} min"
                ),
            }
        )

        invalid = dict(TEST_BATCH)
        invalid["humidity"] = 140
        invalid_response = client.post("/api/predict", json=invalid)
        rows.append(
            {
                "area": "Validation",
                "check": "Reject humidity above 100%",
                "status": "Pass" if invalid_response.status_code == 400 else "Fail",
                "evidence": f"HTTP {invalid_response.status_code}",
            }
        )

        bad_fish = dict(TEST_BATCH)
        bad_fish["fishType"] = "unknown"
        bad_fish_response = client.post("/api/predict", json=bad_fish)
        rows.append(
            {
                "area": "Validation",
                "check": "Reject unsupported fish type",
                "status": "Pass" if bad_fish_response.status_code == 400 else "Fail",
                "evidence": f"HTTP {bad_fish_response.status_code}",
            }
        )

    return rows, prediction_payload


def add_table(document: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)


def add_figure(document: Document, title: str, path: Path, note: str) -> None:
    document.add_heading(title, level=2)
    document.add_picture(str(path), width=Inches(6.45))
    last = document.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph(note)
    paragraph.style = document.styles["Caption"]


def fix_docx_zoom_setting(path: Path) -> None:
    """python-docx can emit a bestFit zoom tag that the schema validator rejects."""

    tmp_path = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/settings.xml":
                text = data.decode("utf-8")
                text = text.replace(
                    '<w:zoom w:val="bestFit"/>',
                    '<w:zoom w:val="bestFit" w:percent="100"/>',
                )
                data = text.encode("utf-8")
            target.writestr(item, data)
    tmp_path.replace(path)


def build_document(
    data: pd.DataFrame,
    evaluation: dict[str, object],
    figures: dict[str, Path],
    decision_matrix: pd.DataFrame,
    app_checks: list[dict[str, str]],
    prediction_payload: dict[str, object],
) -> None:
    metadata = json.loads(METADATA_PATH.read_text(encoding="utf-8"))
    metrics = evaluation["metrics"]

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(23, 32, 28)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI Model and Application Results Report")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Intelligent Solar-Assisted Fish Kiln Monitoring System")
    document.add_paragraph(f"Generated: {date.today().isoformat()}")

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        "This report summarizes the current machine-learning model and application validation results for the fish kiln monitoring prototype. "
        "The model estimates remaining drying time from fish type, weight readings, temperature, humidity, elapsed time, and derived drying measurements."
    )
    document.add_paragraph(
        "Important limitation: the current training dataset is synthetic demonstration data generated by the project pipeline. "
        "It should be presented as a prototype dataset until real experimental drying trials are collected and used for retraining."
    )

    document.add_heading("Dataset and Training Setup", level=1)
    add_table(
        document,
        ["Item", "Value"],
        [
            ["Dataset file", str(DATA_PATH.relative_to(ROOT_DIR))],
            ["Total rows", len(data)],
            ["Fish species", ", ".join(sorted(data["fish_type"].unique()))],
            ["Feature count", len(FEATURE_COLUMNS)],
            ["Target variable", metadata.get("target", "remaining_drying_time")],
            ["Model type", metadata.get("modelType", "Random Forest Regression")],
            ["Train/test split", "80% training, 20% testing, random_state=42"],
        ],
    )

    document.add_heading("Model Evaluation Metrics", level=1)
    add_table(
        document,
        ["Metric", "Value", "Meaning"],
        [
            ["MAE", metrics["mae"], "Average absolute error in minutes"],
            ["RMSE", metrics["rmse"], "Penalizes larger prediction errors"],
            ["R2", metrics["r2"], "Explained variance on the held-out test set"],
        ],
    )
    document.add_paragraph(
        "A standard classification confusion matrix is not directly applicable because the model is a regression model with a continuous target. "
        "To provide a confusion-matrix-like result, the continuous remaining-time output was grouped into operational classes: Ready/complete, Short wait, and Needs drying."
    )

    add_figure(
        document,
        "Correlation Heatmap",
        figures["correlation"],
        "Shows relationships between numeric input and derived drying variables. This helps explain which measurements move together during drying.",
    )
    add_figure(
        document,
        "Actual vs Predicted Plot",
        figures["actual_vs_predicted"],
        "Points close to the diagonal line indicate predictions close to the held-out test target.",
    )
    add_figure(
        document,
        "Residual Distribution",
        figures["residuals"],
        "Residuals centered near zero indicate that the model is not strongly biased toward overprediction or underprediction.",
    )
    add_figure(
        document,
        "Feature Importance",
        figures["feature_importance"],
        "Shows the strongest contributors used by the Random Forest model for remaining-time estimation.",
    )
    add_figure(
        document,
        "Operational Decision Matrix",
        figures["decision_matrix"],
        "A binned regression output matrix used as the practical equivalent of a confusion matrix for the app decision categories.",
    )

    document.add_heading("Decision Matrix Values", level=2)
    add_table(
        document,
        ["Actual / Predicted", *[str(column) for column in decision_matrix.columns]],
        [[str(index), *row.tolist()] for index, row in decision_matrix.iterrows()],
    )

    document.add_heading("Application Validation Matrix", level=1)
    add_table(
        document,
        ["Area", "Check", "Status", "Evidence"],
        [[row["area"], row["check"], row["status"], row["evidence"]] for row in app_checks],
    )

    document.add_heading("Demo Prediction Result", level=1)
    add_table(
        document,
        ["Output", "Value"],
        [
            ["Fish type", prediction_payload.get("fishType", "n/a")],
            ["Estimated moisture content", f"{prediction_payload.get('estimatedMoistureContent', 'n/a')}%"],
            ["Drying progress", f"{prediction_payload.get('dryingProgress', 'n/a')}%"],
            ["Estimated remaining time", f"{prediction_payload.get('estimatedRemainingTime', 'n/a')} minutes"],
            ["Drying status", prediction_payload.get("dryingStatus", "n/a")],
            ["Alerts", ", ".join(prediction_payload.get("alerts", [])) or "None"],
        ],
    )

    document.add_heading("Supervisor Explanation Notes", level=1)
    notes = [
        "The system combines physical drying calculations with an ML model that predicts remaining drying time.",
        "Weight loss is used to estimate current moisture content and drying progress toward the 15% target moisture level.",
        "The current model is trained on synthetic data, so it proves the workflow and app logic but should be retrained with real kiln experiment data.",
        "Regression results are reported with MAE, RMSE, R2, residual plots, feature importance, and actual-vs-predicted plots.",
        "The app was validated through production build checks, API health checks, prediction checks, and bad-input validation checks.",
    ]
    for note in notes:
        document.add_paragraph(note, style=None)

    document.save(REPORT_PATH)
    fix_docx_zoom_setting(REPORT_PATH)


def main() -> None:
    ensure_artifacts()
    data = load_training_frame()
    evaluation = evaluate_model(data)
    app_checks, prediction_payload = run_app_checks()

    figures = {
        "correlation": plot_correlation_heatmap(data),
        "actual_vs_predicted": plot_actual_vs_predicted(
            evaluation["y_test"],
            evaluation["predictions"],
        ),
        "residuals": plot_residuals(evaluation["residuals"]),
        "feature_importance": plot_feature_importance(evaluation["model"]),
    }
    decision_path, decision_matrix = plot_decision_matrix(
        evaluation["y_test"],
        evaluation["predictions"],
    )
    figures["decision_matrix"] = decision_path

    build_document(data, evaluation, figures, decision_matrix, app_checks, prediction_payload)

    print(f"Report written to: {REPORT_PATH}")
    print("Figures written to:", FIGURES_DIR)
    print("Metrics:", evaluation["metrics"])


if __name__ == "__main__":
    main()
